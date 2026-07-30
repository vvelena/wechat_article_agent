import csv
import re
from datetime import date, datetime, timedelta, timezone
from pathlib import Path

from app.core.paths import (
    REPORTS_DIR,
    WECHAT_RESULTS_FILE,
    WEB_RESULTS_FILE,
)
from app.core.result_schema import (
    RESULT_FIELDS,
    ensure_csv_schema,
)


TEMPLATE_FILE = (
    Path(__file__).resolve().parent
    / "templates"
    / "daily_report_template.docx"
)

RESULT_FILES = {
    "wechat": WECHAT_RESULTS_FILE,
    "web": WEB_RESULTS_FILE,
}


def china_today() -> date:
    china_timezone = timezone(timedelta(hours=8))
    return datetime.now(china_timezone).date()


def parse_report_date(report_date: str = "") -> date:
    text = str(report_date or "").strip()
    if not text:
        return china_today()

    normalized = (
        text.replace("年", "-")
        .replace("月", "-")
        .replace("日", "")
        .replace("/", "-")
    )
    try:
        return datetime.strptime(
            normalized,
            "%Y-%m-%d",
        ).date()
    except ValueError as error:
        raise ValueError(
            "report_date 必须是 YYYY-MM-DD 或 YYYY年M月D日。"
        ) from error


def parse_date_from_text(value: str) -> date | None:
    text = str(value or "").strip()
    match = re.search(
        r"(20\d{2})[年\-/.](\d{1,2})[月\-/.](\d{1,2})",
        text,
    )
    if not match:
        return None

    try:
        return date(
            int(match.group(1)),
            int(match.group(2)),
            int(match.group(3)),
        )
    except ValueError:
        return None


def load_result_rows() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []

    for source_kind, csv_file in RESULT_FILES.items():
        if not csv_file.exists() or csv_file.stat().st_size == 0:
            continue

        ensure_csv_schema(csv_file, RESULT_FIELDS)
        with csv_file.open(
            mode="r",
            newline="",
            encoding="utf-8-sig",
        ) as file:
            for row in csv.DictReader(file):
                normalized_row = {
                    key: str(value or "").strip()
                    for key, value in row.items()
                }
                normalized_row["source_kind"] = source_kind
                rows.append(normalized_row)

    return rows


def score_row(row: dict[str, str]) -> float:
    def score(field: str) -> int:
        try:
            return max(0, min(int(row.get(field, "0")), 100))
        except (TypeError, ValueError):
            return 0

    total = (
        score("importance_score") * 0.40
        + score("relevance_score") * 0.25
        + score("quality_score") * 0.20
        + score("source_reliability_score") * 0.15
    )
    if row.get("is_promotional") == "是":
        total -= 15
    return total


def select_daily_rows(
    rows: list[dict[str, str]],
    report_date: date,
    max_items: int = 0,
) -> tuple[list[dict[str, str]], int]:
    max_items = int(max_items)
    if max_items < 0:
        raise ValueError("max_items 不能小于 0。")
    candidates = []
    legacy_fallback_count = 0

    for row in rows:
        collected_date = parse_date_from_text(
            row.get("collected_at", "")
        )
        used_legacy_fallback = collected_date is None
        effective_date = (
            collected_date
            or parse_date_from_text(row.get("publish_time", ""))
        )

        if effective_date != report_date:
            continue
        if row.get("category") == "非量子行业":
            continue
        try:
            relevance_score = int(
                row.get("relevance_score", "0")
            )
        except (TypeError, ValueError):
            relevance_score = 0
        if relevance_score < 60:
            continue

        selected_row = dict(row)
        selected_row["_score"] = str(score_row(row))
        selected_row["_used_legacy_fallback"] = (
            "true" if used_legacy_fallback else "false"
        )
        candidates.append(selected_row)

    candidates.sort(
        key=lambda row: float(row["_score"]),
        reverse=True,
    )
    selected = []
    seen_urls: set[str] = set()
    seen_titles: set[str] = set()

    for row in candidates:
        url = row.get("article_url", "").strip()
        title_key = re.sub(
            r"[\W_]+",
            "",
            row.get("title", "").lower(),
        )
        if url and url in seen_urls:
            continue
        if title_key and title_key in seen_titles:
            continue

        if url:
            seen_urls.add(url)
        if title_key:
            seen_titles.add(title_key)
        selected.append(row)

        if max_items > 0 and len(selected) >= max_items:
            break

    legacy_fallback_count = sum(
        row["_used_legacy_fallback"] == "true"
        for row in selected
    )
    return selected, legacy_fallback_count


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def clear_paragraph(paragraph) -> None:
    paragraph._element.clear_content()


def add_hyperlink(
    paragraph,
    text: str,
    url: str,
) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE

    relationship_id = paragraph.part.relate_to(
        url,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "20")
    properties.extend([color, underline, size])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def render_daily_word(
    selected_rows: list[dict[str, str]],
    report_date: date,
    output_file: Path,
) -> None:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as error:
        raise RuntimeError(
            "缺少 python-docx，请先安装 requirements.txt。"
        ) from error

    if not TEMPLATE_FILE.exists():
        raise FileNotFoundError(
            f"日报模板不存在：{TEMPLATE_FILE}"
        )

    document = Document(TEMPLATE_FILE)
    if len(document.paragraphs) < 3:
        raise ValueError("日报模板结构无效，至少需要三个标题段落。")

    for paragraph in list(document.paragraphs[3:]):
        remove_paragraph(paragraph)

    date_paragraph = document.paragraphs[1]
    clear_paragraph(date_paragraph)
    release_run = date_paragraph.add_run(
        f"发布日期：{report_date.year}年"
        f"{report_date.month}月{report_date.day}日"
    )
    release_run.font.size = Pt(10.5)
    release_run.add_break()
    period_run = date_paragraph.add_run(
        f"统计周期：{report_date.year}年"
        f"{report_date.month}月{report_date.day}日"
    )
    period_run.font.size = Pt(10.5)

    for index, row in enumerate(selected_rows, start=1):
        article_paragraph = document.add_paragraph()
        article_paragraph.paragraph_format.space_after = Pt(5)

        title_run = article_paragraph.add_run(
            f"{index}. {row.get('title', '').strip()}"
        )
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.add_break()

        summary_run = article_paragraph.add_run(
            row.get("summary", "").strip()
        )
        summary_run.bold = True
        summary_run.font.size = Pt(10.5)

        url = row.get("article_url", "").strip()
        if url:
            link_paragraph = document.add_paragraph()
            add_hyperlink(
                link_paragraph,
                "原文链接",
                url,
            )

    output_file.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_file)


def generate_daily_word_report(
    report_date: str = "",
    max_items: int = 0,
) -> dict:
    target_date = parse_report_date(report_date)
    rows = load_result_rows()
    selected_rows, legacy_fallback_count = select_daily_rows(
        rows,
        target_date,
        max_items,
    )

    if not selected_rows:
        return {
            "success": False,
            "status": "no_data",
            "report_date": target_date.isoformat(),
            "message": (
                f"未找到 {target_date.isoformat()} "
                "可进入日报的收集数据，未生成 Word。"
            ),
        }

    base_output = (
        REPORTS_DIR
        / f"量子科技行业资讯日报_{target_date:%Y%m%d}.docx"
    )
    output_file = base_output
    if output_file.exists():
        output_file = REPORTS_DIR / (
            f"量子科技行业资讯日报_{target_date:%Y%m%d}_"
            f"{datetime.now():%H%M%S}.docx"
        )

    render_daily_word(
        selected_rows,
        target_date,
        output_file,
    )

    return {
        "success": True,
        "status": "generated",
        "report_date": target_date.isoformat(),
        "selected_count": len(selected_rows),
        "legacy_fallback_count": legacy_fallback_count,
        "output_file": str(output_file.resolve()),
        "titles": [
            row.get("title", "")
            for row in selected_rows
        ],
    }
